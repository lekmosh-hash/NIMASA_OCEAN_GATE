import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from io import BytesIO

# --- 1. PAGE SETUP ---
st.set_page_config(page_title="NIMASA Ocean-Gate: Forensic Engine", layout="wide")

# Initialize Session State
if 'vessel_db' not in st.session_state:
    st.session_state.vessel_db = []

# --- 2. MOCK DATA GENERATOR (For the Presentation) ---
def seed_data():
    mock_entries = [
        {
            "vessel": "MV THOR CHAICHANA", "imo": "9638501", 
            "eta": datetime.now() + timedelta(hours=2), # Triggering a late penalty
            "sub_time": datetime.now(), "lead_hrs": 2.0,
            "terminal": "ECM Calabar", "sso_phone": "+66-81-234-5678",
            "location": "Approaching Calabar Fairway Buoy", "cargo_desc": "Wheat in Bulk",
            "cargo_qty": 10000.0, "issc": "Valid", "ssp": "Yes", "issuing_govt": "Singapore",
            "sts_history": "None", "history": "1. Tincan NG\n2. Durban SA\n3. Kismayo SO",
            "agent": "Blue Seas Agency", "status": "Violation Detected"
        },
        {
            "vessel": "MT OCEAN PROSPERITY", "imo": "9234567", 
            "eta": datetime.now() + timedelta(days=5), # Compliant
            "sub_time": datetime.now(), "lead_hrs": 120.0,
            "terminal": "Five Star Lagos", "sso_phone": "+44-77-0090-0123",
            "location": "Gulf of Guinea", "cargo_desc": "Crude Oil",
            "cargo_qty": 45000.0, "issc": "Valid", "ssp": "Yes", "issuing_govt": "Panama",
            "sts_history": "STS at Lome Offshore, 2026-03-01", "history": "1. Lome, TG\n2. Luanda, AO",
            "agent": "Global Marine", "status": "Compliant"
        }
    ]
    st.session_state.vessel_db = mock_entries

# --- 3. DOCUMENT GENERATOR ---
def generate_forensic_brief(v):
    doc = Document()
    doc.add_heading('NIMASA ISPS-SHIP UNIT: FORENSIC AUDIT', 0)
    doc.add_heading('Statutory Compliance Summary', level=1)
    doc.add_paragraph(f"Vessel: {v['vessel']} | IMO: {v['imo']}")
    doc.add_paragraph(f"Vetting Status: {'VIOLATION (Part VIII/8)' if v['lead_hrs'] < 72 else 'COMPLIANT'}")
    doc.add_paragraph(f"Lead Time: {v['lead_hrs']:.1f} hours")
    doc.add_paragraph(f"Terminal Joint Liability: {v['terminal']} (Applicable if berthed)")
    
    doc.add_heading('Security Particulars', level=1)
    doc.add_paragraph(f"SSP on Board: {v['ssp']} (Issued by {v['issuing_govt']})")
    doc.add_paragraph(f"Last 10 Ports: {v['history']}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. SIDEBAR NAVIGATION ---
st.sidebar.title("🛡️ Ocean-Gate Control")
mode = st.sidebar.radio("Navigate To:", ["1. Shipowner/Agent Entry Portal", "2. NIMASA Command Center"])

st.sidebar.divider()
if st.sidebar.button("🚀 Seed Presentation Mock Data"):
    seed_data()
    st.rerun()

if st.sidebar.button("🧹 Clear Database"):
    st.session_state.vessel_db = []
    st.rerun()

# =========================================================
# INTERFACE 1: SHIPOWNER FORM
# =========================================================
if mode == "1. Shipowner/Agent Entry Portal":
    st.header("📤 Pre-Arrival Information (PAI) Statutory Filing")
    st.markdown("#### Direct Gateway for SOLAS XI-2/9 & NIMASA ISPS Regulations 2014")

    with st.form("forensic_pai_form", clear_on_submit=True):
        st.subheader("Section A: Identity & Contact")
        c1, c2, c3 = st.columns(3)
        with c1:
            v_name = st.text_input("Vessel Name")
            imo = st.text_input("IMO Number")
        with c2:
            agent_name = st.text_input("Local Agent Name")
            sso_phone = st.text_input("SSO Satellite Contact")
        with c3:
            v_pos = st.text_input("Current Position")
            terminal = st.selectbox("Assign Terminal", ["ECM Calabar", "Five Star Lagos", "Onne", "PTML"])

        st.subheader("Section B: Cargo & Certification")
        c4, c5, c6 = st.columns(3)
        with c4:
            eta = st.datetime_input("Expected Arrival Time (ETA)")
            cargo_desc = st.text_input("Cargo Description")
        with c5:
            issc = st.selectbox("ISSC Status", ["Valid", "Expired", "Interim"])
            ssp = st.radio("SSP on Board?", ["Yes", "No"])
        with c6:
            issuing_govt = st.text_input("Issuing Administration")
            cargo_qty = st.number_input("Cargo Quantity (MT)")

        st.subheader("Section C: Forensic History")
        sts = st.text_area("STS History (Proposed/Previous)")
        ports = st.text_area("Last 10 Ports of Call (Mandatory)")

        if st.form_submit_button("SUBMIT STATUTORY FILING"):
            sub_time = datetime.now()
            lead = (eta - sub_time).total_seconds() / 3600
            st.session_state.vessel_db.append({
                "vessel": v_name, "imo": imo, "eta": eta, "sub_time": sub_time,
                "lead_hrs": lead, "terminal": terminal, "sso_phone": sso_phone,
                "location": v_pos, "cargo_desc": cargo_desc, "cargo_qty": cargo_qty,
                "issc": issc, "ssp": ssp, "issuing_govt": issuing_govt,
                "sts_history": sts, "history": ports, "agent": agent_name
            })
            st.success("PAI Logged. Access via Command Center.")

# =========================================================
# INTERFACE 2: COMMAND CENTER
# =========================================================
elif mode == "2. NIMASA Command Center":
    st.header("🛡️ ISPS Forensic Command Dashboard")
    
    if not st.session_state.vessel_db:
        st.warning("No filings pending. Use the sidebar to Seed Mock Data or the Portal to enter manually.")
    else:
        for v in st.session_state.vessel_db:
            with st.expander(f"AUDIT: {v['vessel']} (Terminal: {v['terminal']})", expanded=True):
                st.write(f"*Agent:* {v['agent']} | *Position:* {v['location']}")
                
                c1, c2, c3 = st.columns(3)
                with c1:
                    st.metric("PAI Lead Time", f"{v['lead_hrs']:.1f} hrs")
                    if v['lead_hrs'] < 72:
                        st.error("🚩 72HR BREACH: Part VIII/8")
                        st.write("Penalty: N300,000 (Shipowner)")
                        st.write(f"Joint Fine: N300,000 ({v['terminal']})")
                with c2:
                    st.write("*Security Vetting*")
                    st.write(f"SSP: {v['ssp']} ({v['issuing_govt']})")
                    if "Kismayo" in v['history']: st.error("🚩 HIGH RISK PORT (Kismayo)")
                with c3:
                    st.write("*Cargo Oversight*")
                    st.write(f"{v['cargo_desc']} ({v['cargo_qty']:,} MT)")
                    if v['sts_history'] != "None": st.warning("🔍 STS Activity Flag")

                st.download_button(f"Download Audit Brief: {v['vessel']}", generate_forensic_brief(v), f"NIMASA_Audit_{v['vessel']}.docx")